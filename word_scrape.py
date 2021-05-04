import re
import json
import docx
from docx2python import docx2python
from docx.enum.text import WD_TAB_ALIGNMENT

#Create Class
class WordScrape:
    def parse_word(self, file_location):

        #Define variables
        combo_filing = []
        exhibits = []
        emerging_growth = []
        period = []
        items = []
        eight_k = {}

        # CREATE WORD DOCUMENT OBJECT
        try:
            doc = docx.Document(file_location)
            docs = docx2python(file_location, extract_image=True)
        except:
            print('Only ".DOCX" files are excepted.')
            return None
        # CREATE PARAGRAPH OBJECT
        paragraphs = doc.paragraphs

        # LOOP THROUGH LINES OF PARAGRAPH. FIND CENTER ALIGNED TEXT THAT EQUAL "FORM 8-K" AND HOUSE IT IN A VARIABLE
        for line in paragraphs:

            if line.alignment == WD_TAB_ALIGNMENT.CENTER and (form_type := line.text.strip().title()) == 'Form 8-K':
                eight_k['FORM_TYPE'] = form_type

        # IDENTIFYING THE DATE OF REPORT
        for line in paragraphs:
            if line.alignment == WD_TAB_ALIGNMENT.CENTER and 'Date Of Report' in line.text.strip().title():
                if not (date_of_report := line.text.strip()):
                    print("Missing Date Of Report")
                else:
                    period.append(date_of_report.split(":")[1:][0])
        eight_k['PERIOD'] = period

        # IDENTIFY FILER
        search_text = [i.text for i in paragraphs if
                       i.alignment == WD_TAB_ALIGNMENT.CENTER and re.findall('^\(Exact', i.text)]
        search_list = [i.text for i in paragraphs if i.alignment == WD_TAB_ALIGNMENT.CENTER]
        search_results = search_list.index(search_text[0])
        try:
            filer = search_list[search_results - 1]
        except:
            filer = ['Not Located']

        eight_k['FILER'] = filer

        # LOCATE COMBO FILINGS
        combo_filing_list = [i for x in docs.body for i in x if i[0][0] == '☐']
        combo_filing.append([[i[0][0], i[1][0]] for i in combo_filing_list])
        eight_k['COMBO FILING'] = combo_filing

        # IDENTIFYING EMERGING GROUPS
        emerging_growth_list = [i for x in docs.body for i in x[0][0]]
        emerging_growth_strip = [i.strip() for i in emerging_growth_list if 'Emerging Growth' in i.title()]
        for line in emerging_growth_strip:
            check_box = re.search('.$', line.strip())
            emerging_text = line.strip().split(f'{check_box.group()}')[0]
            emerging_growth.append(f'{check_box.group()} {emerging_text.strip()}')

        eight_k['EMERGING GROWTH'] = emerging_growth

        # IDENTIFYING "ITEMS"
        paragraph = [i for x in paragraphs for i in x.runs]
        items.insert(0, [i.text.strip() for i in paragraph if 'Item' in i.text])
        eight_k['ITEMS'] = items

        # LOCATE EXHIBITS BY LOOPING THROUGH THE TABLES IN THE DOCUMENT AND LOCATING THE KEY IDENTIFIER "☐"
        for num in range(len(doc.tables)):
            for row in doc.tables[num].rows:
                if 'Exhibit' in row.cells[0].text:
                    for exhibit_row in doc.tables[num].rows:
                        value_check = re.findall("\d.+", exhibit_row.cells[0].text)
                        if not value_check:
                            continue
                        else:
                            exhibits.append(exhibit_row.cells[0].text)
                if 'Item' in row.cells[0].text:
                    items.append(row.cells[0].text)

        eight_k['EXHIBITS'] = exhibits


        # JSON PRINT STATEMENT
        print(json.dumps(eight_k, indent=3, ensure_ascii=False), '\n')
